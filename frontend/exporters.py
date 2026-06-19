from __future__ import annotations

import re
from io import BytesIO


def _bold_runs(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        run.bold = part.startswith("**") and part.endswith("**")


def to_docx(md: str) -> bytes:
    from docx import Document

    doc = Document()
    for raw in md.splitlines():
        s = raw.rstrip()
        if not s.strip():
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.lstrip().startswith(("- ", "* ")):
            _bold_runs(doc.add_paragraph(style="List Bullet"), s.lstrip()[2:].strip())
        elif s.startswith("|"):  # keep tables readable as plain rows
            doc.add_paragraph(s)
        else:
            _bold_runs(doc.add_paragraph(), s)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _latin(text: str) -> str:
    # fpdf core fonts are latin-1; swap common unicode and drop the rest.
    swaps = {"—": "-", "–": "-", "•": "-", "✓": "[x]", "⚠": "(!)", "“": '"', "”": '"', "’": "'", "→": "->"}
    for a, b in swaps.items():
        text = text.replace(a, b)
    return text.encode("latin-1", "replace").decode("latin-1")


def to_pdf(md: str) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()

    def write(text: str, size: int, style: str = "", lh: int = 6) -> None:
        pdf.set_font("Helvetica", style, size)
        pdf.multi_cell(0, lh, _latin(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    for raw in md.splitlines():
        s = raw.rstrip()
        if not s.strip():
            pdf.ln(2)
            continue
        if s.startswith("### "):
            write(s[4:].strip(), 12, "B", 7)
        elif s.startswith("## "):
            write(s[3:].strip(), 14, "B", 8)
        elif s.startswith("# "):
            write(s[2:].strip(), 18, "B", 9)
        else:
            write(re.sub(r"\*\*(.*?)\*\*", r"\1", s), 11, "", 6)
    return bytes(pdf.output())
